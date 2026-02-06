"""Convert the integration framework markdown to a formatted .docx document."""

import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def add_code_block(doc, text, font_size=8):
    """Add a formatted code block paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    # Set background shading on the paragraph
    shading = p.paragraph_format.element.get_or_add_pPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): "F0F0F5",
    })
    shading.append(shd)
    return p


def add_table(doc, headers, rows):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
        # Shade header
        shading = cell._element.get_or_add_tcPr()
        shd = shading.makeelement(qn("w:shd"), {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): "2B3A67",
        })
        shading.append(shd)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[1 + row_idx].cells[col_idx]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()  # spacing after table


def parse_table(lines, start_idx):
    """Parse a markdown table starting at start_idx. Returns (headers, rows, end_idx)."""
    # Header line
    header_line = lines[start_idx].strip()
    headers = [h.strip() for h in header_line.split("|") if h.strip()]

    # Skip separator line
    separator_idx = start_idx + 1

    # Data rows
    rows = []
    idx = separator_idx + 1
    while idx < len(lines):
        line = lines[idx].strip()
        if not line.startswith("|"):
            break
        row = [c.strip() for c in line.split("|") if c.strip()]
        # Pad row if needed
        while len(row) < len(headers):
            row.append("")
        rows.append(row[:len(headers)])
        idx += 1

    return headers, rows, idx


def build_document():
    md_path = Path(__file__).parent / "LLM_GAME_ENGINE_INTEGRATION_FRAMEWORK.md"
    md_text = md_path.read_text()
    lines = md_text.split("\n")

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Title
    title = doc.add_heading("LLM Agent + Python Game Engine Integration Framework", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "A Dual-Representation Architecture for ARC-AGI-3 Domain Modeling\n"
        "Integrating the LearningAgent's Structured Natural Language Pipeline "
        "with an Executable Python Symbolic Model\n\n"
        "February 2026"
    )
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.italic = True

    doc.add_page_break()

    # Process markdown line by line
    i = 0
    in_code_block = False
    code_buffer = []

    while i < len(lines):
        line = lines[i]

        # Skip the first few lines (title, subtitle, date) — already handled
        if i < 8:
            i += 1
            continue

        # Code blocks
        if line.strip().startswith("```"):
            if in_code_block:
                # End code block
                add_code_block(doc, "\n".join(code_buffer))
                code_buffer = []
                in_code_block = False
            else:
                # Start code block
                in_code_block = True
                code_buffer = []
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        stripped = line.strip()

        # Skip horizontal rules
        if stripped == "---" or stripped == "═══":
            i += 1
            continue

        # Skip table of contents links
        if re.match(r"^\d+\.\s+\[", stripped):
            i += 1
            continue

        # Tables
        if stripped.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            # Check if next line is separator
            next_line = lines[i + 1].strip()
            if re.match(r"^\|[\s\-:|]+\|", next_line):
                headers, rows, end_idx = parse_table(lines, i)
                add_table(doc, headers, rows)
                i = end_idx
                continue

        # Headings
        if stripped.startswith("## ") and not stripped.startswith("### "):
            heading_text = stripped.lstrip("# ").strip()
            # Remove markdown anchor links
            heading_text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", heading_text)
            doc.add_heading(heading_text, level=1)
            i += 1
            continue

        if stripped.startswith("### ") and not stripped.startswith("#### "):
            heading_text = stripped.lstrip("# ").strip()
            heading_text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", heading_text)
            doc.add_heading(heading_text, level=2)
            i += 1
            continue

        if stripped.startswith("#### "):
            heading_text = stripped.lstrip("# ").strip()
            heading_text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", heading_text)
            doc.add_heading(heading_text, level=3)
            i += 1
            continue

        # Empty lines
        if not stripped:
            i += 1
            continue

        # Bullet points
        if stripped.startswith("- ") or stripped.startswith("* "):
            bullet_text = stripped[2:]
            p = doc.add_paragraph(style="List Bullet")
            # Handle inline code and bold
            _add_formatted_text(p, bullet_text)
            i += 1
            continue

        # Numbered lists
        match = re.match(r"^(\d+)\.\s+(.+)", stripped)
        if match:
            list_text = match.group(2)
            p = doc.add_paragraph(style="List Number")
            _add_formatted_text(p, list_text)
            i += 1
            continue

        # Blockquote
        if stripped.startswith("> "):
            quote_text = stripped[2:]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(quote_text)
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # Regular paragraph — handle inline formatting
        p = doc.add_paragraph()
        _add_formatted_text(p, stripped)
        i += 1

    # Save
    output_path = Path(__file__).parent / "LLM_GAME_ENGINE_INTEGRATION_FRAMEWORK.docx"
    doc.save(str(output_path))
    print(f"Saved to {output_path}")
    return output_path


def _add_formatted_text(paragraph, text):
    """Add text to a paragraph with inline formatting (bold, italic, code)."""
    # Split on inline code, bold, and italic markers
    # Process tokens: **bold**, *italic*, `code`
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", text)

    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x80, 0x20, 0x20)
        else:
            paragraph.add_run(part)


if __name__ == "__main__":
    build_document()
